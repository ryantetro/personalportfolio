#!/usr/bin/env python3
"""Build public/Ryan_Tetro_Resume.docx from the same content as resume/Ryan_Tetro_Resume.md."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "public" / "Ryan_Tetro_Resume.docx"

LINKEDIN_URL = "https://www.linkedin.com/in/ryan-tetro/"

EDUCATION = [
    (
        "Apr 2026",
        "Master of Science, Information Systems, Brigham Young University - Marriott School of Business",
        "Provo",
    ),
    (
        "Apr 2025",
        "Bachelor of Science, Information Systems, Brigham Young University - Marriott School of Business",
        "Provo",
    ),
    (
        "Apr 2026",
        "Sandbox Fellowship — Startup Incubator Program, Brigham Young University",
        "Provo",
    ),
]

JOBS = [
    (
        "Sept 2025 — Present",
        "Founder & Software Engineer, PostGame AI",
        "Provo",
        [
            "Architected and built an AI-powered coaching platform enabling coaches to send audio feedback and athletes to receive AI-generated insights and performance summaries",
            "Developed full-stack application using Next.js, React, Supabase, and PostgreSQL supporting authentication, data storage, and scalable user workflows",
            "Implemented AI pipelines using OpenAI APIs to transcribe, analyze, and generate structured feedback from coach audio recordings",
            "Designed backend infrastructure and APIs to handle audio processing, transcript storage, and real-time AI analysis",
            "Led product development from concept to deployment, including system architecture, feature design, and implementation",
        ],
    ),
    (
        "Jan 2025 — Present",
        "Software Engineer, Soar (AI Technology Startup)",
        "Provo",
        [
            "Built a production AI-powered lead generation platform using Next.js, AWS Lambda, DynamoDB, and Pinecone vector database",
            "Implemented semantic search pipelines using OpenAI embeddings and vector similarity search to improve lead targeting",
            "Developed scalable serverless backend infrastructure and REST APIs to support real-time data processing and AI workflows",
            "Designed database architecture and backend systems supporting production deployment and multi-user workflows",
        ],
    ),
    (
        "Feb 2024 — Jan 2025",
        "Data Engineer, Soar (AI technology start-up)",
        "Provo",
        [
            "Develop data AI models that use gen AI to take real-time information, existing information and provide AI analysis, insights and actions to users",
            "Engineer, in Python, ML modules that customers use daily",
            "Work with key customers to create AI modules that can be repeatably used inside organizations to train, inform and change user behavior",
        ],
    ),
    (
        "Dec 2023 — Present",
        "Investment Analyst, Utah Innovation Fund (Investment Fund)",
        "Salt Lake City",
        [
            "Analyze and evaluate investment opportunities, successfully identifying high-potential startups focused in deep tech and life sciences sectors",
            "Develop financial models and forecasts for potential investments, providing risk assessments and market analysis that informed the fund's investment strategies",
        ],
    ),
    (
        "Oct 2022 — Oct 2023",
        "Senior Data Associate, Brandless (Healthcare technology start-up)",
        "Lindon",
        [
            "Create data dashboards in Sigma to help the business make critical day to day decisions",
            "Work with the data and insights team to structure and represent information to the CEO and executives in the company",
            "Work on strategies for how data can be used to continue to grow the business",
        ],
    ),
    (
        "Mar 2021 — Oct 2022",
        "M&A Associate, Brandless (Healthcare technology start-up)",
        None,
        [
            "Source potential acquisitions targets",
            "Qualify opportunities by presenting to customers and entrepreneurs",
            "Consistently met quota for lead, qualification and opportunities for the company resulting in helping drive the growth of the company 1400%",
        ],
    ),
    (
        "Nov 2020 — Mar 2021",
        "Business Development Representative, GlassFrogg (Healthcare technology start-up)",
        "Provo",
        [
            "Worked with a cross-functional team to identify target customer, determine the customer journey and drive",
            "Increased revenue 100% monthly through cold calls, lead generation and company presentations",
            "Provided customer insight for the product development team, customer support team and management to help customer acquisition through relationships improve customer experience",
        ],
    ),
    (
        "Oct 2020 — Present",
        "Co-founder & Partner, SurfMongers (Wake surf start-up)",
        None,
        [
            "Developed a light up wake surf product line to create a visible surf experience for surf enthusiast, this included sourcing materials, developing prototypes, testing products",
            "Built a go-to-market strategy for the product launch, fundraising and influencer campaigns",
            "Created a crowdfunding campaign, that is just about to start, to fund the first product releases",
            "Structured a company with partners and investors to support the launch and growth",
            "Worked with key business mentors to develop and build the company from the ground",
        ],
    ),
    (
        "May 2019 — Present",
        "Youth Leader and Soccer Coach, BYU Soccer Youth Program Coach (University)",
        "Provo",
        [
            "Worked with youth teams through soccer skills and leadership activities",
            "Led my teams to two championship outcomes over the course of the summer",
            "Helped youth learn about soccer, life skills and how to work with others",
        ],
    ),
    (
        "May 2018 — Aug 2018",
        "Tech Entrepreneur Apprentice, Album VC (Tech Apprenticeship program)",
        "Lehi",
        [
            "Worked with over 20 tech entrepreneurs to learn customer validation, market sizing, fundraising, customer service, leadership, team development, product development and testing",
            "Worked on projects to help build and validation technology ideas that could develop into a start up",
            "Worked with a venture capital company in Silicon Slopes to develop skills to help build tech companies",
        ],
    ),
]

LEADERSHIP = [
    (
        "2021 — Present",
        "BYU Soccer Team — Player",
        None,
        None,
        [
            "Won national championships 5 years in a row, 2021, 2022, 2023, 2024, and 2025",
        ],
    ),
    (
        "2015 — 2019",
        "Team Captain (LaRoca Soccer Team & AF High School Team)",
        "Provo, UT",
        "Team Captain",
        [
            "Led our team to Regionals",
            "Drove player and team performance in the DPL soccer league",
            "Led 40 high school students to 3 years in regionals and top 5 in state",
        ],
    ),
    (
        "2019 — 2020",
        "The Church of Jesus Christ of Latter-day Saints",
        "Quito, Ecuador",
        "Volunteer Representative",
        [
            "Selected to serve in leadership position responsible for the training and performance of volunteers",
            "Tripled the statistical success in assigned area for 6 consecutive months through exceptional time management",
        ],
    ),
]

SKILLS = [
    ("Spanish: Proficient working ability", "Expert"),
    ("Scrum Master Certified", "Expert"),
    ("Project management", "Expert"),
    ("Agile Development training", "Expert"),
    ("User Interface Design", "Expert"),
    ("Figma", "Expert"),
    ("Customer Research", "Expert"),
    ("Unified Modeling Language (UML)", "Expert"),
    ("Proficient in SQL", "Expert"),
    ("MySQL", "Expert"),
    ("Data Exploratory analysis in Python", "Expert"),
    ("Creation and Deployment of dashboards in Python", "Expert"),
    ("JavaScript/Node/Express", "Expert"),
    ("PostgreSQL", "Expert"),
    ("Full-Stack Programming using ASP.net Core MVC / C#", "Expert"),
    ("API & SPA Development", "Expert"),
    ("JQuery", "Expert"),
    ("React", "Expert"),
    ("SQLite", "Expert"),
    ("Linux Proficiency", "Expert"),
    ("Encryption/Decryption Techniques", "Expert"),
    ("Web server security", "Expert"),
    ("Digital Signatures", "Expert"),
    ("Use of SSH", "Expert"),
    ("password Managers", "Expert"),
]


def set_normal_style(document: Document) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 1.15


def set_margins(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_section_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(10)


def add_date_title_line(document: Document, date_str: str, title_str: str) -> None:
    p = document.add_paragraph()
    ts = p.paragraph_format.tab_stops
    ts.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    r0 = p.add_run(date_str)
    r0.bold = True
    p.add_run("\t")
    r1 = p.add_run(title_str)
    r1.bold = True
    p.paragraph_format.space_after = Pt(0)


def add_contact(document: Document) -> None:
    t = document.add_paragraph()
    r = t.add_run("Ryan Tetro")
    r.bold = True
    r.font.size = Pt(16)
    t.paragraph_format.space_after = Pt(6)

    p1 = document.add_paragraph("Cedar Hills, Utah, United States, 8013610618,")
    p2 = document.add_paragraph()
    link = p2.add_hyperlink(LINKEDIN_URL, "LinkedIn")
    # python-docx Hyperlink: use add_run with style or oxlm - simpler: mailto + plain linkedin
    _ = link  # noqa: F841 — Hyperlink helper below
    p2.clear()
    run_mail = p2.add_run("ryantetro@gmail.com")
    run_mail.font.color.rgb = None
    p2 = document.add_paragraph()
    add_hyperlink(p2, "LinkedIn", LINKEDIN_URL)


def add_hyperlink(paragraph, text: str, url: str):
    """Append a hyperlink to paragraph (external URL)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_education_block(document: Document) -> None:
    add_section_heading(document, "Education")
    for date_s, degree, place in EDUCATION:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        d = p.add_run(date_s)
        d.bold = True
        p.add_run(" — ")
        p.add_run(degree)
        lp = document.add_paragraph(place)
        lp.paragraph_format.left_indent = Inches(0)
        lf = lp.runs[0].font if lp.runs else None
        if lf is not None:
            lf.italic = True
        lp.paragraph_format.space_after = Pt(6)


def add_job_block(document: Document) -> None:
    add_section_heading(document, "Employment history")
    for date_s, title, location, bullets in JOBS:
        add_date_title_line(document, date_s, title)
        if location:
            loc = document.add_paragraph(location)
            for r in loc.runs:
                r.italic = True
            loc.paragraph_format.space_after = Pt(2)
        for b in bullets:
            document.add_paragraph(b, style="List Bullet")
        document.add_paragraph().paragraph_format.space_after = Pt(4)


def add_leadership_block(document: Document) -> None:
    add_section_heading(document, "Leadership & service")
    for date_s, title, place, role_line, bullets in LEADERSHIP:
        add_date_title_line(document, date_s, title)
        if place:
            pl = document.add_paragraph(place)
            for r in pl.runs:
                r.italic = True
        if role_line:
            rl = document.add_paragraph(role_line)
            for r in rl.runs:
                r.bold = True
        for b in bullets:
            document.add_paragraph(b, style="List Bullet")
        document.add_paragraph().paragraph_format.space_after = Pt(4)


def add_skills_table(document: Document) -> None:
    add_section_heading(document, "Awards & skills")
    table = document.add_table(rows=len(SKILLS), cols=2)
    table.autofit = False
    for i, (skill, level) in enumerate(SKILLS):
        row = table.rows[i]
        row.cells[0].text = skill
        row.cells[1].text = level
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    if c == row.cells[1]:
                        r.bold = True
    document.add_paragraph()


def add_languages(document: Document) -> None:
    add_section_heading(document, "Languages")
    document.add_paragraph("Spanish: Proficient working ability", style="List Bullet")


def fix_contact_email(document: Document) -> None:
    """Replace botched email paragraph with mailto hyperlink."""
    # Email as plain text paragraph (second body paragraph after name)
    paras = document.paragraphs
    # Index 0: name, 1: address line — insert email before LinkedIn
    # Current add_contact order: name, line1, p2 botched, we need: name, line1, email, linkedin
    pass


def build() -> Path:
    document = Document()
    set_margins(document)
    set_normal_style(document)

    # Contact: name, address, email paragraph, linkedin
    t = document.add_paragraph()
    r = t.add_run("Ryan Tetro")
    r.bold = True
    r.font.size = Pt(16)
    t.paragraph_format.space_after = Pt(6)
    document.add_paragraph("Cedar Hills, Utah, United States, 8013610618,")
    pe = document.add_paragraph()
    add_hyperlink(pe, "ryantetro@gmail.com", "mailto:ryantetro@gmail.com")
    pl = document.add_paragraph()
    add_hyperlink(pl, "LinkedIn", LINKEDIN_URL)

    add_education_block(document)
    add_job_block(document)
    add_leadership_block(document)
    add_skills_table(document)
    add_languages(document)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
